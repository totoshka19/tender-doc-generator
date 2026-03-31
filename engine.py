import json
import os
from datetime import date
from pathlib import Path

import yaml
from docx import Document


# ---------------------------------------------------------------------------
# 5.1  Загрузка данных
# ---------------------------------------------------------------------------

MONTHS_RU = [
    "января", "февраля", "марта", "апреля", "мая", "июня",
    "июля", "августа", "сентября", "октября", "ноября", "декабря",
]


def load_data(profile_path: str, tender_path: str, calc_path: str) -> dict:
    """Читает три JSON и возвращает единый контекст с ключами profile/tender/calc/system."""
    with open(profile_path, encoding="utf-8") as f:
        profile = json.load(f)
    with open(tender_path, encoding="utf-8") as f:
        tender = json.load(f)
    with open(calc_path, encoding="utf-8") as f:
        calc = json.load(f)

    today = date.today()
    return {
        "profile": profile,
        "tender": tender,
        "calc": calc,
        "system": {
            "date_long": f"«{today.day}» {MONTHS_RU[today.month - 1]} {today.year} года",
            "outgoing_number": f"{today.strftime('%d%m')}/1",
        },
    }


# ---------------------------------------------------------------------------
# 5.2  Форматирование чисел
# ---------------------------------------------------------------------------

def format_money(value) -> str:
    """432600.0 → '432 600,00'  (пробел — разделитель тысяч, запятая — дробная часть)."""
    formatted = f"{float(value):,.2f}"   # '432,600.00'
    return formatted.replace(",", " ").replace(".", ",")


# ---------------------------------------------------------------------------
# 5.3  Резолвер путей
# ---------------------------------------------------------------------------

def resolve(path: str, context: dict) -> str:
    """
    Принимает путь вида 'profile.company.inn' или 'calc.subtotal_wo_vat | money'
    и возвращает строковое значение из контекста.
    Если ключ не найден — возвращает пустую строку, не выбрасывает ошибку.
    """
    parts = [p.strip() for p in path.split("|")]
    key_path = parts[0]
    apply_money = "money" in parts[1:]

    value = context
    for key in key_path.split("."):
        if not isinstance(value, dict):
            return ""
        value = value.get(key)
        if value is None:
            return ""

    if apply_money:
        return format_money(value)
    return str(value)


# ---------------------------------------------------------------------------
# 5.4  Безопасная замена плейсхолдера в параграфе
# ---------------------------------------------------------------------------

def replace_in_paragraph(paragraph, placeholder: str, value: str) -> None:
    """
    Склеивает текст всех runs параграфа, ищет плейсхолдер в склеенной строке.
    При нахождении: записывает результат в первый run, остальные обнуляет.
    Форматирование первого run сохраняется.
    """
    full_text = "".join(run.text for run in paragraph.runs)
    if placeholder not in full_text:
        return
    new_text = full_text.replace(placeholder, value)
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""


# ---------------------------------------------------------------------------
# 5.5  Обход документа и применение простых замен
# ---------------------------------------------------------------------------

def fill_document(doc, replacements: dict) -> None:
    """Применяет словарь замен ко всем параграфам и ячейкам таблиц документа."""
    for paragraph in doc.paragraphs:
        for placeholder, value in replacements.items():
            replace_in_paragraph(paragraph, placeholder, value)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for placeholder, value in replacements.items():
                        replace_in_paragraph(paragraph, placeholder, value)


# ---------------------------------------------------------------------------
# 5.6  Заполнение таблицы позиций (документ 03)
# ---------------------------------------------------------------------------

def fill_items_table(doc, table_config: dict, context: dict) -> None:
    """
    Заполняет таблицу позиций в документе 03.
    Для каждой строки строит локальный контекст с calc_item и tender_item,
    затем записывает значения по индексам столбцов из конфига.
    """
    table = doc.tables[table_config["table_index"]]
    header_rows = table_config.get("header_rows", 1)
    columns: dict = table_config["columns"]  # {col_index: path}

    calc_items = context["calc"].get("items", [])
    tender_items = context["tender"].get("items", [])

    for row_idx, row in enumerate(table.rows[header_rows:]):
        if row_idx >= len(calc_items):
            break

        row_context = {
            **context,
            "calc_item": calc_items[row_idx],
            "tender_item": tender_items[row_idx] if row_idx < len(tender_items) else {},
        }

        for col_idx_str, path in columns.items():
            col_idx = int(col_idx_str)
            if col_idx >= len(row.cells):
                continue
            value = resolve(path, row_context)
            para = row.cells[col_idx].paragraphs[0]
            if para.runs:
                para.runs[0].text = value
                for run in para.runs[1:]:
                    run.text = ""
            else:
                para.add_run(value)


# ---------------------------------------------------------------------------
# 5.7  Главная функция генерации одного документа
# ---------------------------------------------------------------------------

def generate_document(mapping_path: str, context: dict) -> None:
    """
    Читает YAML-маппинг, открывает шаблон, выполняет замены,
    заполняет таблицу позиций (если указана) и сохраняет результат.
    """
    with open(mapping_path, encoding="utf-8") as f:
        mapping = yaml.safe_load(f)

    replacements_config: dict = mapping.get("replacements", {})
    replacements = {
        placeholder: resolve(path, context)
        for placeholder, path in replacements_config.items()
    }

    doc = Document(mapping["template"])
    fill_document(doc, replacements)

    if "items_table" in mapping:
        fill_items_table(doc, mapping["items_table"], context)

    out_path = Path(mapping["output"])
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    print(f"Создан: {out_path}")
