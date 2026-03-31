import json
import re
import sys

from docx import Document


def _find_paragraph(texts: list[str], keyword: str) -> str:
    for text in texts:
        if keyword in text:
            return text
    return ""


def _after_label(text: str, label: str) -> str:
    """Возвращает текст после метки вида 'Метка: значение', без финальной точки."""
    idx = text.find(label)
    if idx == -1:
        return ""
    after = text[idx + len(label):]
    return after.lstrip(": ").strip().rstrip(".")


def _strip_numbering(text: str) -> str:
    """Убирает нумерацию вида '1.2. ' из начала строки."""
    return re.sub(r"^\d+\.\d+\.\s*", "", text)


def _parse_float(s: str) -> float | None:
    try:
        return float(s.replace(" ", "").replace(",", "."))
    except (ValueError, AttributeError):
        return None


def _parse_int(s: str) -> int | None:
    try:
        return int(s.replace(" ", ""))
    except (ValueError, AttributeError):
        return None


def _parse_purchase_line(text: str) -> dict:
    """Разбирает строку 'Номер закупки: X    Лот: Y    Код лота: Z'."""
    result = {"purchase_number": None, "lot_number": None, "lot_code": None}
    m = re.search(r"Номер закупки:\s*(\S+)", text)
    if m:
        result["purchase_number"] = m.group(1).strip()
    m = re.search(r"(?<!\S)Лот:\s*(.+?)(?:\s{2,}|$)", text)
    if m:
        result["lot_number"] = m.group(1).strip()
    m = re.search(r"Код лота:\s*(\S+)", text)
    if m:
        result["lot_code"] = m.group(1).strip()
    return result


def _parse_items_table(table) -> list:
    """Разбирает таблицу позиций, пропуская строку заголовка."""
    items = []
    for row in table.rows[1:]:
        cells = [c.text.strip() for c in row.cells]
        if not any(cells):
            continue
        items.append({
            "line_no": _parse_int(cells[0]),
            "customer_name_code": cells[1] or None,
            "article": cells[2] or None,
            "name": cells[3] or None,
            "qty": _parse_int(cells[4]),
            "unit": cells[5] or None,
            "nmc_unit_price": _parse_float(cells[6]),
            "required_delivery_date": cells[7] or None,
            "customer_org": None,
            "basis": None,
        })
    return items


def extract(docx_path: str) -> dict:
    doc = Document(docx_path)
    texts = [p.text.strip() for p in doc.paragraphs]

    purchase_line = _find_paragraph(texts, "Номер закупки:")
    purchase_info = _parse_purchase_line(purchase_line)

    customer_para = _find_paragraph(texts, "Заказчик:")
    customer_full_name = _after_label(customer_para, "Заказчик") or None

    contact_para = _find_paragraph(texts, "e-mail:")
    customer_email = None
    m = re.search(r"e-mail:\s*(\S+)", contact_para)
    if m:
        customer_email = m.group(1).strip()

    subject_para = _find_paragraph(texts, "Предмет закупки:")
    subject = _after_label(subject_para, "Предмет закупки") or None

    deadline_para = _find_paragraph(texts, "Срок подачи предложений:")
    bid_deadline = _after_label(deadline_para, "Срок подачи предложений") or None

    place_para = _find_paragraph(texts, "Место поставки:")
    delivery_place = _after_label(place_para, "Место поставки") or None

    term_para = _find_paragraph(texts, "Срок поставки:")
    delivery_term_text = _strip_numbering(term_para).strip() or None

    payment_para = _find_paragraph(texts, "Оплата")
    payment_term_text = _strip_numbering(payment_para).strip() or None

    warranty_para = _find_paragraph(texts, "Гарантийный срок")
    warranty_term_text = warranty_para or None

    items = []
    for table in doc.tables:
        headers = [c.text.strip() for c in table.rows[0].cells]
        if "Наименование" in headers:
            items = _parse_items_table(table)
            break

    return {
        "purchase_number": purchase_info["purchase_number"],
        "lot_number": purchase_info["lot_number"],
        "lot_code": purchase_info["lot_code"],
        "subject": subject,
        "bid_deadline": bid_deadline,
        "offer_validity_days": None,
        "contract_number": None,
        "currency": None,
        "customer": {
            "full_name": customer_full_name,
            "short_name": None,
            "legal_address": None,
            "postal_address": None,
            "email": customer_email,
            "inn": None,
            "kpp": None,
            "ogrn": None,
            "bank": {
                "name": None,
                "account": None,
                "correspondent_account": None,
                "bik": None,
            },
            "signatory": {
                "position": None,
                "name": None,
            },
        },
        "delivery": {
            "place": delivery_place,
            "basis": None,
            "start_text": None,
            "end_text": None,
            "term_text": delivery_term_text,
        },
        "payment": {
            "term_text": payment_term_text,
        },
        "warranty": {
            "term_text": warranty_term_text,
        },
        "items": items,
    }


def main():
    if len(sys.argv) < 2:
        print("Использование: python extractor.py <путь_к_docx> [выходной_json]")
        sys.exit(1)

    docx_path = sys.argv[1]
    out_path = sys.argv[2] if len(sys.argv) > 2 else "data/tender.json"

    result = extract(docx_path)
    with open(out_path, "w", encoding="utf-8") as f:
        json.dump(result, f, ensure_ascii=False, indent=2)

    print(f"Извлечено: {out_path}")


if __name__ == "__main__":
    main()
